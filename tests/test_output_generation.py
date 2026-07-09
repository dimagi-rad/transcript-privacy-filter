from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document

from opf_app.models import ParsedItem, RedactionResult
from opf_app.outputs import (
    generate_docx_output,
    generate_v2_docx_output,
    package_successful_outputs,
    package_v2_successful_outputs,
)


def test_docx_output_contains_title_metadata_and_body(tmp_path: Path) -> None:
    result = _result(
        source_name="example-input.csv",
        redacted_text="[2026-05-11 10:00:00+00:00] User: <PRIVATE_PERSON>",
        session_id="session-1",
    )

    output = generate_docx_output(result, tmp_path)

    assert output.success is True
    assert output.path is not None
    paragraphs = _paragraphs(output.path)
    assert "Redacted Transcript" in paragraphs
    assert "Source: example-input.csv" in paragraphs
    assert "Chat date: 2026-05-11" in paragraphs
    assert "User identifier: user-1" in paragraphs
    assert "Session ID: session-1" in paragraphs
    assert "[2026-05-11 10:00:00+00:00] USER: <PRIVATE_PERSON>" in paragraphs


def test_docx_output_formats_transcript_speaker_runs(tmp_path: Path) -> None:
    result = _result(
        source_name="example-input.csv",
        redacted_text=(
            "[2026-05-11 10:00:00+00:00] User: Hello\n"
            "[<PRIVATE_DATE>] Chatbot: Hi\n"
            "[00:00] S1: Opening line\n"
            "Plain continuation"
        ),
    )

    output = generate_docx_output(result, tmp_path)

    assert output.path is not None
    document = Document(output.path)
    user_paragraph = _paragraph_with_text(
        document,
        "[2026-05-11 10:00:00+00:00] USER: Hello",
    )
    chatbot_paragraph = _paragraph_with_text(document, "[<PRIVATE_DATE>] CHATBOT: Hi")
    speaker_paragraph = _paragraph_with_text(document, "[00:00] S1: Opening line")
    plain_paragraph = _paragraph_with_text(document, "Plain continuation")

    assert [run.text for run in user_paragraph.runs] == [
        "[2026-05-11 10:00:00+00:00] ",
        "USER:",
        " Hello",
    ]
    assert user_paragraph.runs[0].bold is not True
    assert user_paragraph.runs[1].bold is True
    assert user_paragraph.runs[2].bold is not True
    assert chatbot_paragraph.runs[1].text == "CHATBOT:"
    assert chatbot_paragraph.runs[1].bold is True
    assert speaker_paragraph.runs[1].text == "S1:"
    assert speaker_paragraph.runs[1].bold is True
    assert all(run.bold is not True for run in plain_paragraph.runs)


def test_zip_contains_successful_outputs_and_handles_filename_collisions(
    tmp_path: Path,
) -> None:
    first = _result(source_name="source-a.txt", redacted_text="first", session_id="a")
    second = _result(source_name="source-b.txt", redacted_text="second", session_id="b")

    package = package_successful_outputs((first, second), tmp_path)

    assert package.zip_path.is_file()
    assert len(package.successful_outputs) == 2
    filenames = [output.filename for output in package.successful_outputs]
    assert len(set(filenames)) == 2
    assert filenames[0] == "redacted-transcript-2026-05-11-user-1.docx"
    assert filenames[1].startswith("redacted-transcript-2026-05-11-user-1-")
    with zipfile.ZipFile(package.zip_path) as archive:
        assert sorted(archive.namelist()) == sorted(filenames)


def test_failed_items_are_excluded_from_zip(tmp_path: Path) -> None:
    success = _result(source_name="source-a.txt", redacted_text="redacted")
    failed = _result(
        source_name="source-b.txt",
        redacted_text="original",
        success=False,
        errors=("redaction failed",),
    )

    package = package_successful_outputs((success, failed), tmp_path)

    assert len(package.outputs) == 2
    assert len(package.successful_outputs) == 1
    assert len(package.failed_outputs) == 1
    assert package.failed_outputs[0].path is None
    with zipfile.ZipFile(package.zip_path) as archive:
        assert archive.namelist() == [success.output_filename]


def test_v2_docx_output_uses_minimal_metadata_without_source_name(
    tmp_path: Path,
) -> None:
    result = _result(
        source_name="sensitive-source-name.txt",
        redacted_text="[2026-05-11 10:00:00+00:00] User: <PRIVATE_PERSON>",
        session_id="session-1",
    )

    output = generate_v2_docx_output(result, tmp_path)

    assert output.success is True
    assert output.path is not None
    paragraphs = _paragraphs(output.path)
    assert "Redacted Transcript" in paragraphs
    assert "Redaction workflow: OpenAI Responses API v2" in paragraphs
    assert "Chat date: 2026-05-11" in paragraphs
    assert "User identifier: user-1" in paragraphs
    assert "Session ID: session-1" in paragraphs
    assert "Source: sensitive-source-name.txt" not in paragraphs
    assert "[2026-05-11 10:00:00+00:00] USER: <PRIVATE_PERSON>" in paragraphs


def test_v2_docx_output_removes_only_xml_invalid_characters(
    tmp_path: Path,
) -> None:
    result = _result(
        source_name="sensitive-source-name.txt",
        redacted_text=(
            "[00:00] User: café\x00before\x08after\x0bvalid\ttab\ufffe"
        ),
        session_id="session\x00-1",
    )

    output = generate_v2_docx_output(result, tmp_path)

    assert output.success is True
    assert output.path is not None
    paragraphs = _paragraphs(output.path)
    assert "Session ID: session-1" in paragraphs
    assert "[00:00] USER: cafébeforeaftervalid\ttab" in paragraphs


def test_docx_output_removes_xml_invalid_characters_from_source_metadata(
    tmp_path: Path,
) -> None:
    result = _result(
        source_name="example\x00-input.csv",
        redacted_text="Plain readable body",
    )

    output = generate_docx_output(result, tmp_path)

    assert output.path is not None
    assert "Source: example-input.csv" in _paragraphs(output.path)


def test_v2_zip_contains_only_successful_outputs_with_existing_filename_rules(
    tmp_path: Path,
) -> None:
    success = _result(source_name="source-a.txt", redacted_text="redacted")
    failed = _result(
        source_name="source-b.txt",
        redacted_text="original",
        success=False,
        errors=("v2_redaction_failed:rate_limit",),
    )

    package = package_v2_successful_outputs((success, failed), tmp_path)

    assert len(package.outputs) == 2
    assert len(package.successful_outputs) == 1
    assert len(package.failed_outputs) == 1
    assert package.successful_outputs[0].filename == success.output_filename
    assert package.failed_outputs[0].path is None
    with zipfile.ZipFile(package.zip_path) as archive:
        assert archive.namelist() == [success.output_filename]


def _result(
    *,
    source_name: str,
    redacted_text: str,
    session_id: str | None = None,
    success: bool = True,
    errors: tuple[str, ...] = (),
) -> RedactionResult:
    item = ParsedItem(
        item_name=source_name,
        source_name=source_name,
        source_type="document",
        chat_date="2026-05-11",
        user_identifier="user-1",
        body_text="original",
        output_filename="redacted-transcript-2026-05-11-user-1.docx",
        session_id=session_id,
    )
    return RedactionResult(
        item=item,
        output_filename=item.output_filename,
        redacted_text=redacted_text,
        success=success,
        errors=errors,
    )


def _paragraphs(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs]


def _paragraph_with_text(document: Document, text: str):
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise AssertionError(f"Paragraph not found: {text}")
