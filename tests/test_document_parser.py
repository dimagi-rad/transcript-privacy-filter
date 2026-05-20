from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfWriter

from tests.fixtures import write_sample_docx
from opf_app.documents import normalize_transcript_row, parse_document


def test_sample_docx_extracts_transcript_body_and_metadata(tmp_path: Path) -> None:
    path = write_sample_docx(tmp_path / "P695 Interview 2026_01_07.docx")

    item = parse_document(path)

    assert item.chat_date == "2026-01-07"
    assert item.user_identifier == "P695"
    assert "[00:00] S1:" in item.body_text
    assert "Transcription details" not in item.body_text
    assert "Input sound file" not in item.body_text
    assert item.output_filename == "redacted-transcript-2026-01-07-P695.docx"


def test_normalize_transcript_row_variants() -> None:
    assert (
        normalize_transcript_row("S1: 00:00 | Okay.")
        == "[00:00] S1: Okay."
    )
    assert (
        normalize_transcript_row("s2 | 01:02:03 | Fine.")
        == "[01:02:03] S2: Fine."
    )
    assert normalize_transcript_row("Date | 2026-01-07") is None


def test_docx_metadata_rows_are_excluded(tmp_path: Path) -> None:
    path = tmp_path / "P123 Interview 2026-01-07.docx"
    document = Document()
    table = document.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "Transcription details"
    table.rows[0].cells[1].text = "metadata"
    table.rows[1].cells[0].text = "Input sound file"
    table.rows[1].cells[1].text = "recording.wav"
    table.rows[2].cells[0].text = "S1:"
    table.rows[2].cells[1].text = "00:00"
    table.rows[2].cells[2].text = "Hello."
    document.save(path)

    item = parse_document(path)

    assert item.body_text == "[00:00] S1: Hello."


def test_txt_parser_reads_utf8_text(tmp_path: Path) -> None:
    path = tmp_path / "Notes 2026-01-07.txt"
    path.write_text("Plain UTF-8 cafe text", encoding="utf-8")

    item = parse_document(path)

    assert item.parse_status == "parsed"
    assert item.body_text == "Plain UTF-8 cafe text"
    assert item.output_filename.endswith(".docx")


def test_empty_txt_returns_warning_item(tmp_path: Path) -> None:
    path = tmp_path / "Empty 2026-01-07.txt"
    path.write_text("", encoding="utf-8")

    item = parse_document(path)

    assert item.parse_status == "warning"
    assert item.warnings == ("Source document is empty.",)
    assert item.body_text == ""


def test_text_pdf_extracts_embedded_text(tmp_path: Path) -> None:
    path = tmp_path / "Report 2026-01-07.pdf"
    _write_text_pdf(path, "Embedded PDF text")

    item = parse_document(path)

    assert item.parse_status == "parsed"
    assert "Embedded PDF text" in item.body_text


def test_empty_pdf_returns_ocr_not_supported_error(tmp_path: Path) -> None:
    path = tmp_path / "Scan 2026-01-07.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as stream:
        writer.write(stream)

    item = parse_document(path)

    assert item.parse_status == "error"
    assert item.body_text == ""
    assert "OCR/image-only PDFs are not supported" in item.errors[0]


def test_doc_converter_success_is_used(tmp_path: Path) -> None:
    doc_path = tmp_path / "P456 Interview 2026-02-03.doc"
    doc_path.write_bytes(b"legacy doc placeholder")
    converted_path = tmp_path / "converted.docx"
    document = Document()
    document.add_paragraph("Converted legacy document body")
    document.save(converted_path)

    item = parse_document(doc_path, doc_converter=lambda _path: converted_path)

    assert item.parse_status == "parsed"
    assert item.chat_date == "2026-02-03"
    assert item.user_identifier == "P456"
    assert item.body_text == "Converted legacy document body"


def test_doc_missing_converter_returns_setup_error(tmp_path: Path, monkeypatch) -> None:
    doc_path = tmp_path / "Legacy 2026-02-03.doc"
    doc_path.write_bytes(b"legacy doc placeholder")
    monkeypatch.setattr("opf_app.documents.shutil.which", lambda _name: None)

    item = parse_document(doc_path)

    assert item.parse_status == "error"
    assert "LibreOffice/soffice is required" in item.errors[0]


def _write_text_pdf(path: Path, text: str) -> None:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(f'BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET'.encode('ascii'))} >>\n"
        f"stream\nBT /F1 12 Tf 72 720 Td ({escaped}) Tj ET\nendstream".encode(
            "ascii"
        ),
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, payload in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(payload)
        pdf.extend(b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(pdf))
