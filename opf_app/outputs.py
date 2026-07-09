from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
import zipfile

from docx import Document

from .filenames import OUTPUT_EXTENSION, stable_suffix
from .models import RedactionResult


_TRANSCRIPT_LINE_RE = re.compile(
    r"^(?P<timestamp>\[[^\]]+\]\s+)"
    r"(?P<speaker>[A-Za-z][A-Za-z0-9 _.-]{0,40}?)"
    r"(?P<separator>:\s*)"
    r"(?P<body>.*)$"
)
V2_WORKFLOW_LABEL = "OpenAI Responses API v2"


@dataclass(frozen=True)
class GeneratedOutput:
    """Status for one attempted generated output file."""

    source_name: str
    filename: str
    path: Path | None
    success: bool
    errors: tuple[str, ...] = field(default_factory=tuple)


@dataclass(frozen=True)
class OutputPackage:
    """Generated output files plus the zip archive containing successes."""

    zip_path: Path
    outputs: tuple[GeneratedOutput, ...]

    @property
    def successful_outputs(self) -> tuple[GeneratedOutput, ...]:
        return tuple(output for output in self.outputs if output.success)

    @property
    def failed_outputs(self) -> tuple[GeneratedOutput, ...]:
        return tuple(output for output in self.outputs if not output.success)


def generate_docx_output(
    result: RedactionResult,
    output_dir: str | Path,
    *,
    filename: str | None = None,
    workflow_label: str | None = None,
    include_source: bool = True,
) -> GeneratedOutput:
    """Generate one plain DOCX output file for a successful redaction result."""
    if not result.success:
        return GeneratedOutput(
            source_name=result.item.source_name,
            filename=filename or result.output_filename,
            path=None,
            success=False,
            errors=result.errors or ("Redaction did not complete successfully.",),
        )

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_filename = filename or result.output_filename
    output_path = output_dir / output_filename
    _write_redacted_docx(
        result,
        output_path,
        workflow_label=workflow_label,
        include_source=include_source,
    )
    return GeneratedOutput(
        source_name=result.item.source_name,
        filename=output_filename,
        path=output_path,
        success=True,
    )


def generate_v2_docx_output(
    result: RedactionResult,
    output_dir: str | Path,
    *,
    filename: str | None = None,
) -> GeneratedOutput:
    """Generate one v2 DOCX output with minimal privacy-safe metadata."""
    return generate_docx_output(
        result,
        output_dir,
        filename=filename,
        workflow_label=V2_WORKFLOW_LABEL,
        include_source=False,
    )


def package_successful_outputs(
    results: list[RedactionResult] | tuple[RedactionResult, ...],
    output_dir: str | Path,
    *,
    zip_name: str = "redacted-transcripts.zip",
) -> OutputPackage:
    """Write successful DOCX outputs and package them into a zip archive."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    used_filenames: set[str] = set()
    outputs: list[GeneratedOutput] = []
    for result in results:
        filename = _unique_filename(result, used_filenames)
        used_filenames.add(filename)
        outputs.append(generate_docx_output(result, output_dir, filename=filename))

    zip_path = output_dir / zip_name
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for output in outputs:
            if output.success and output.path is not None:
                archive.write(output.path, arcname=output.filename)

    return OutputPackage(zip_path=zip_path, outputs=tuple(outputs))


def package_v2_successful_outputs(
    results: list[RedactionResult] | tuple[RedactionResult, ...],
    output_dir: str | Path,
    *,
    zip_name: str = "redacted-transcripts.zip",
) -> OutputPackage:
    """Write successful v2 DOCX outputs and package them into a zip archive."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    used_filenames: set[str] = set()
    outputs: list[GeneratedOutput] = []
    for result in results:
        filename = _unique_filename(result, used_filenames)
        used_filenames.add(filename)
        outputs.append(generate_v2_docx_output(result, output_dir, filename=filename))

    zip_path = output_dir / zip_name
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for output in outputs:
            if output.success and output.path is not None:
                archive.write(output.path, arcname=output.filename)

    return OutputPackage(zip_path=zip_path, outputs=tuple(outputs))


def _write_redacted_docx(
    result: RedactionResult,
    output_path: Path,
    *,
    workflow_label: str | None,
    include_source: bool,
) -> None:
    document = Document()
    document.add_heading("Redacted Transcript", level=1)
    if workflow_label:
        document.add_paragraph(f"Redaction workflow: {workflow_label}")
    if include_source:
        document.add_paragraph(f"Source: {result.item.source_name}")
    document.add_paragraph(f"Chat date: {result.item.chat_date}")
    document.add_paragraph(f"User identifier: {result.item.user_identifier}")
    if result.item.session_id:
        document.add_paragraph(f"Session ID: {result.item.session_id}")

    document.add_paragraph("")
    for line in result.redacted_text.splitlines() or [""]:
        _add_body_paragraph(document, line)

    document.save(output_path)


def _add_body_paragraph(document: Document, line: str) -> None:
    match = _TRANSCRIPT_LINE_RE.match(line)
    if match is None:
        document.add_paragraph(line)
        return

    paragraph = document.add_paragraph()
    paragraph.add_run(match.group("timestamp"))
    speaker_run = paragraph.add_run(f"{match.group('speaker').upper()}:")
    speaker_run.bold = True
    paragraph.add_run(match.group("separator")[1:] + match.group("body"))


def _unique_filename(result: RedactionResult, used_filenames: set[str]) -> str:
    filename = result.output_filename
    if filename not in used_filenames:
        return filename

    path = Path(filename)
    extension = path.suffix or OUTPUT_EXTENSION
    stem = path.stem
    discriminator = (
        result.item.session_id
        or result.item.source_path
        or result.item.item_name
        or result.item.source_name
    )
    candidate = f"{stem}-{stable_suffix(discriminator)}{extension}"
    counter = 1
    while candidate in used_filenames:
        counter += 1
        candidate = f"{stem}-{stable_suffix(f'{discriminator}:{counter}')}{extension}"
    return candidate
