from __future__ import annotations

from collections.abc import Callable
from pathlib import Path
import re
import shutil
import subprocess
import tempfile

from .filenames import document_output_filename, infer_chat_date_from_stem, infer_user_identifier_from_stem
from .models import ParsedItem, ParseStatus


SUPPORTED_DOCUMENT_EXTENSIONS = frozenset({".txt", ".docx", ".doc", ".pdf"})
DocConverter = Callable[[Path], Path]

_TRANSCRIPT_ROW_PATTERN = re.compile(
    r"^\s*(?P<speaker>[A-Za-z]\d+)\s*:?\s*(?:\|\s*)?"
    r"(?P<timestamp>\d{1,2}:\d{2}(?::\d{2})?)\s*(?:\|\s*|[-]\s+|\s+)"
    r"(?P<text>.+?)\s*$"
)
_METADATA_KEYS = ("transcription details", "date", "input sound file")


class DocumentParseError(ValueError):
    """Raised when a document cannot produce redaction-ready body text."""


def parse_document(
    path: str | Path,
    *,
    doc_converter: DocConverter | None = None,
) -> ParsedItem:
    """Parse one supported document into a redaction-ready item."""
    source_path = Path(path)
    chat_date = infer_chat_date_from_stem(source_path.stem)
    user_identifier = infer_user_identifier_from_stem(source_path.stem)
    output_filename = document_output_filename(source_path.name)

    try:
        body_text = _extract_document_body(source_path, doc_converter=doc_converter)
    except DocumentParseError as exc:
        return _build_document_item(
            source_path,
            chat_date=chat_date,
            user_identifier=user_identifier,
            output_filename=output_filename,
            body_text="",
            parse_status="error",
            errors=(str(exc),),
        )

    warnings: tuple[str, ...] = ()
    parse_status: ParseStatus = "parsed"
    if not body_text:
        warnings = ("Source document is empty.",)
        parse_status = "warning"

    return _build_document_item(
        source_path,
        chat_date=chat_date,
        user_identifier=user_identifier,
        output_filename=output_filename,
        body_text=body_text,
        parse_status=parse_status,
        warnings=warnings,
    )


def parse_document_folder(path: str | Path) -> tuple[ParsedItem, ...]:
    """Parse all supported non-zipped documents in a folder."""
    folder = Path(path)
    if not folder.is_dir():
        raise DocumentParseError(f"Document folder does not exist: {folder}")
    return tuple(
        parse_document(candidate)
        for candidate in sorted(folder.iterdir())
        if candidate.is_file() and candidate.suffix.lower() in SUPPORTED_DOCUMENT_EXTENSIONS
    )


def normalize_transcript_row(line: str) -> str | None:
    """Normalize transcript rows like 'S1: 00:00 | text'."""
    match = _TRANSCRIPT_ROW_PATTERN.match(line)
    if not match:
        return None
    speaker = match.group("speaker").upper()
    timestamp = match.group("timestamp")
    text = match.group("text").strip()
    if not text:
        return None
    return f"[{timestamp}] {speaker}: {text}"


def _extract_document_body(
    path: Path,
    *,
    doc_converter: DocConverter | None = None,
) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        return _extract_docx_body(path)
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".doc":
        return _extract_doc_body(path, doc_converter=doc_converter)
    raise DocumentParseError(f"Unsupported document type: {suffix or '<none>'}")


def _extract_docx_body(path: Path) -> str:
    from docx import Document

    document = Document(path)
    raw_lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            raw_lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            ]
            if cells:
                raw_lines.append(" | ".join(cells))

    return _body_from_raw_lines(raw_lines)


def _body_from_raw_lines(lines: list[str]) -> str:
    transcript_lines: list[str] = []
    fallback_lines: list[str] = []

    for line in lines:
        if _is_metadata_line(line):
            continue
        normalized = normalize_transcript_row(line)
        if normalized is not None:
            transcript_lines.append(normalized)
        elif line.strip():
            fallback_lines.append(line.strip())

    return "\n".join(transcript_lines if transcript_lines else fallback_lines)


def _is_metadata_line(line: str) -> bool:
    normalized = line.strip().lower().strip(":")
    return any(
        normalized == key
        or normalized.startswith(f"{key}:")
        or normalized.startswith(f"{key} |")
        for key in _METADATA_KEYS
    )


def _extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    page_text = [
        text.strip()
        for page in reader.pages
        for text in [page.extract_text() or ""]
        if text.strip()
    ]
    if not page_text:
        raise DocumentParseError(
            "No embedded text found in PDF; OCR/image-only PDFs are not supported."
        )
    return "\n".join(page_text)


def _extract_doc_body(
    path: Path,
    *,
    doc_converter: DocConverter | None = None,
) -> str:
    if doc_converter is not None:
        return _extract_docx_body(doc_converter(path))

    with tempfile.TemporaryDirectory() as tmpdir:
        converted_path = _convert_doc_with_soffice(path, Path(tmpdir))
        return _extract_docx_body(converted_path)


def _convert_doc_with_soffice(path: Path, output_dir: Path) -> Path:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        raise DocumentParseError(
            "LibreOffice/soffice is required to parse .doc files. "
            "Install LibreOffice or convert this file to .docx."
        )

    completed = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(output_dir),
            str(path),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    if completed.returncode != 0:
        raise DocumentParseError(
            "LibreOffice failed to convert .doc file to .docx."
        )

    converted_path = output_dir / f"{path.stem}.docx"
    if not converted_path.is_file():
        raise DocumentParseError(
            "LibreOffice conversion did not produce a .docx file."
        )
    return converted_path


def _build_document_item(
    source_path: Path,
    *,
    chat_date: str,
    user_identifier: str,
    output_filename: str,
    body_text: str,
    parse_status: ParseStatus,
    warnings: tuple[str, ...] = (),
    errors: tuple[str, ...] = (),
) -> ParsedItem:
    line_count = len([line for line in body_text.splitlines() if line.strip()])
    return ParsedItem(
        item_name=source_path.name,
        source_name=source_path.name,
        source_type="document",
        chat_date=chat_date,
        user_identifier=user_identifier,
        body_text=body_text,
        output_filename=output_filename,
        parse_status=parse_status,
        source_path=str(source_path),
        line_count=line_count,
        character_count=len(body_text),
        warnings=warnings,
        errors=errors,
    )
